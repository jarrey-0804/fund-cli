"""
Word报告生成器.

使用 python-docx 生成 Word 格式的基金分析报告。
"""
import tempfile
from pathlib import Path
from typing import Any

from fund_cli.core.reporter import Reporter


class DocxReporter(Reporter):
    """Word报告生成器."""

    def generate(self, fund_code: str, metrics: dict[str, Any], nav_data: Any = None, benchmark_data: Any = None, **kwargs) -> str:  # type: ignore[override]
        """生成Word文档内容（返回临时文件路径）."""
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor  # Cm, Inches 未使用
        except ImportError as exc:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx") from exc

        doc = Document()

        # 标题
        title = doc.add_heading(f'{fund_code} 基金分析报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 日期
        from datetime import date
        p = doc.add_paragraph(f'报告日期: {date.today().strftime("%Y年%m月%d日")}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 投资摘要
        doc.add_heading('投资摘要', level=1)
        summary = doc.add_paragraph()
        total_return = metrics.get('total_return', 0)
        run = summary.add_run(f'总收益率: {total_return:.2%}')
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if total_return > 0 else RGBColor(0xE7, 0x4C, 0x3C)
        summary.add_run(f'  |  夏普比率: {metrics.get("sharpe_ratio", "N/A")}')
        summary.add_run(f'  |  最大回撤: {metrics.get("max_drawdown", "N/A")}')

        # 核心绩效指标
        doc.add_heading('核心绩效指标', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '值'

        key_metrics = [
            ("总收益率", metrics.get("total_return", "N/A")),
            ("年化收益率", metrics.get("annualized_return", "N/A")),
            ("波动率", metrics.get("volatility", "N/A")),
            ("夏普比率", metrics.get("sharpe_ratio", "N/A")),
            ("最大回撤", metrics.get("max_drawdown", "N/A")),
            ("Alpha", metrics.get("alpha", "N/A")),
            ("Beta", metrics.get("beta", "N/A")),
        ]
        for name, value in key_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = f"{value:.4f}" if isinstance(value, (int, float)) else str(value)

        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph('本报告由 Fund CLI v3.1 自动生成，仅供参考，不构成投资建议。')
        footer.style.font.size = Pt(9)

        # 保存到临时文件
        temp_path = Path(tempfile.gettempdir()) / f"{fund_code}_report.docx"
        doc.save(str(temp_path))
        return str(temp_path)

    def save(self, content: str, output_path: str) -> None:
        """保存Word文档."""
        import shutil
        shutil.copy2(content, output_path)

    def get_formats(self) -> list[str]:
        return ["docx"]
